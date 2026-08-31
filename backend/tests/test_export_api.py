import json
from io import BytesIO
from zipfile import ZipFile

from docx import Document
from fastapi.testclient import TestClient
from lxml import etree

from app.main import app
from app.services.auth_store import AuthStore
from app.services.docx_modifier import modify_docx_inplace


client = TestClient(app)
W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _build_docx_bytes() -> bytes:
    document = Document()
    styled_paragraph = document.add_paragraph("合同份数：一式两份。")
    styled_paragraph.style = "List Paragraph"
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "签订地点：未约定。"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_english_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("(i) Representation by Counsel. Each of the parties acknowledges that it has entered into this Agreement based upon its independent judgment.")
    document.add_paragraph("(j) Counterparts. This Agreement may be executed in one or more counterparts, each of which shall be considered an original instrument.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_docx_with_same_header_text() -> bytes:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "合同份数：一式两份。"
    document.add_paragraph("合同份数：一式两份。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _read_docx_xml(docx_bytes: bytes, path: str) -> etree._Element:
    with ZipFile(BytesIO(docx_bytes), "r") as archive:
        return etree.fromstring(archive.read(path))


def _paragraph_texts_from_xml(root: etree._Element) -> list[str]:
    paragraphs = []
    for paragraph in root.findall(".//w:p", W_NS):
        parts = [
            node.text or ""
            for node in paragraph.xpath(".//w:t | .//w:delText", namespaces=W_NS)
        ]
        joined = "".join(parts).strip()
        if joined:
            paragraphs.append(joined)
    return paragraphs


def test_export_returns_modified_docx() -> None:
    modifications = [
        {
            "original": "合同份数：一式两份。",
            "modified": "合同份数：一式四份，甲乙双方各执两份。",
        },
        {
            "original": "签订地点：未约定。",
            "modified": "签订地点：上海市浦东新区。",
        },
    ]

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"modifications": json.dumps(modifications, ensure_ascii=False)},
    )

    assert response.status_code == 200
    assert response.headers["content-disposition"] == 'attachment; filename="reviewed_contract.docx"'

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    settings_xml = _read_docx_xml(response.content, "word/settings.xml")

    assert document_xml.find(".//w:ins", W_NS) is not None
    assert document_xml.find(".//w:del", W_NS) is not None
    assert settings_xml.find(".//w:trackRevisions", W_NS) is not None


def test_tracked_export_uses_the_modification_author() -> None:
    result = modify_docx_inplace(
        _build_docx_bytes(),
        [{
            "original": "合同份数：一式两份。",
            "modified": "合同份数：一式四份。",
            "author_display_name": "甲同事",
        }],
    )

    document_xml = _read_docx_xml(result.content, "word/document.xml")
    authors = document_xml.xpath(".//w:ins/@w:author | .//w:del/@w:author", namespaces=W_NS)
    assert authors == ["甲同事", "甲同事"]


def test_export_uses_the_authenticated_user_not_a_client_supplied_author(monkeypatch, tmp_path) -> None:
    auth_db = tmp_path / "auth.sqlite3"
    monkeypatch.setenv("AUTH_DB", str(auth_db))
    monkeypatch.setenv("REVIEW_JOB_WORKER_ENABLED", "false")
    AuthStore(auth_db).create_user("alice", "甲同事", "correct-password")

    with TestClient(app) as authenticated_client:
        assert authenticated_client.post(
            "/api/auth/login", json={"username": "alice", "password": "correct-password"}
        ).status_code == 200
        response = authenticated_client.post(
            "/api/export",
            files={"file": ("contract.docx", _build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"modifications": json.dumps([{
                "original": "合同份数：一式两份。",
                "modified": "合同份数：一式四份。",
                "author_display_name": "伪造作者",
            }], ensure_ascii=False)},
        )

    assert response.status_code == 200
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    assert document_xml.xpath(".//w:ins/@w:author", namespaces=W_NS) == ["甲同事"]


def test_tracked_export_preserves_multiple_revisions_in_one_paragraph() -> None:
    document = Document()
    document.add_paragraph("付款应在签约后支付，验收标准由乙方确定。")
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {"original": "签约后支付", "modified": "验收合格并收到发票后 30 日内支付"},
                    {"original": "由乙方确定", "modified": "由甲方书面确认"},
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    assert len(document_xml.findall(".//w:ins", W_NS)) == 2
    assert len(document_xml.findall(".//w:del", W_NS)) == 2
    inserted_text = "".join(document_xml.xpath(".//w:ins//w:t/text()", namespaces=W_NS))
    deleted_text = "".join(document_xml.xpath(".//w:del//w:delText/text()", namespaces=W_NS))
    assert "验收合格并收到发票后 30 日内支付" in inserted_text
    assert "由甲方书面确认" in inserted_text
    assert "签约后支付" in deleted_text
    assert "由乙方确定" in deleted_text


def test_tracked_export_supports_whole_paragraph_deletion() -> None:
    document = Document()
    document.add_paragraph("第一段保留。")
    document.add_paragraph("第二段需要删除。")
    document.add_paragraph("第三段保留。")
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"modifications": json.dumps([{"original": "第二段需要删除。", "modified": ""}], ensure_ascii=False)},
    )

    assert response.status_code == 200
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    settings_xml = _read_docx_xml(response.content, "word/settings.xml")
    assert document_xml.xpath(".//w:p[w:pPr/w:rPr/w:del]", namespaces=W_NS)
    deleted_text = "".join(document_xml.xpath(".//w:del//w:delText/text()", namespaces=W_NS))
    assert "第二段需要删除。" in deleted_text
    assert settings_xml.find(".//w:trackRevisions", W_NS) is not None


def test_final_export_removes_a_whole_deleted_paragraph() -> None:
    document = Document()
    document.add_paragraph("第一段保留。")
    document.add_paragraph("第二段需要删除。")
    document.add_paragraph("第三段保留。")
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps([{"original": "第二段需要删除。", "modified": ""}], ensure_ascii=False),
            "export_mode": "final",
        },
    )

    assert response.status_code == 200
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    text = "".join(document_xml.xpath(".//w:t/text()", namespaces=W_NS))
    assert "第二段需要删除。" not in text
    assert "第一段保留。" in text
    assert "第三段保留。" in text


def test_tracked_export_supports_precise_sentence_deletion() -> None:
    document = Document()
    document.add_paragraph("付款安排：验收后支付，保留付款期限。")
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [{
                    "original": "验收后支付",
                    "modified": "",
                    "paragraph_context": "付款安排：验收后支付，保留付款期限。",
                }],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    deleted_text = "".join(document_xml.xpath(".//w:del//w:delText/text()", namespaces=W_NS))
    visible_text = "".join(document_xml.xpath(".//w:t/text()", namespaces=W_NS))
    assert deleted_text == "验收后支付"
    assert "付款安排：" in visible_text
    assert "保留付款期限。" in visible_text
    assert "验收后支付" not in visible_text


def test_final_export_supports_precise_sentence_deletion() -> None:
    document = Document()
    document.add_paragraph("付款安排：验收后支付，保留付款期限。")
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps([{"original": "验收后支付", "modified": "", "paragraph_context": "付款安排：验收后支付，保留付款期限。"}], ensure_ascii=False),
            "export_mode": "final",
        },
    )

    assert response.status_code == 200
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    text = "".join(document_xml.xpath(".//w:t/text()", namespaces=W_NS))
    assert text == "付款安排：，保留付款期限。"


def test_export_rejects_fuzzy_replacement_to_preserve_precise_location() -> None:
    modifications = [
        {
            "original": "合同份数一式两份",
            "modified": "合同份数：一式三份，甲乙双方各执一份，存档一份。",
        }
    ]

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"modifications": json.dumps(modifications, ensure_ascii=False)},
    )

    assert response.status_code == 400
    assert "could not be located exactly" in response.json()["detail"]


def test_export_skips_unlocated_suggestion_when_other_revision_is_precise() -> None:
    """Pending or stale suggestions cannot block an otherwise valid review export."""
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {
                        "original": "合同份数：一式两份。",
                        "modified": "合同份数：一式四份，甲乙双方各执两份。",
                    },
                    {
                        "original": "不存在的原文定位",
                        "modified": "这条建议仍保留在审核页面，不能写入 Word。",
                    },
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200
    assert response.headers["x-review-requested-modifications"] == "2"
    assert response.headers["x-review-applied-modifications"] == "1"
    assert response.headers["x-review-skipped-modifications"] == "1"
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    inserted_text = "".join(document_xml.xpath(".//w:ins//w:t/text()", namespaces=W_NS))
    assert "合同份数：一式四份" in inserted_text


def test_export_rejects_repeated_original_to_preserve_precise_location() -> None:
    document = Document()
    document.add_paragraph("付款应在验收后支付。")
    document.add_paragraph("付款应在验收后支付。")
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [{"original": "付款应在验收后支付。", "modified": "付款应在验收合格后 30 日内支付。"}],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 400
    assert "could not be located exactly" in response.json()["detail"]


def test_export_allows_repeated_quote_after_user_confirms_its_paragraph() -> None:
    document = Document()
    document.add_paragraph("付款条件：付款应在验收后支付。")
    document.add_paragraph("补充约定：付款应在验收后支付。")
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [{
                    "original": "付款应在验收后支付。",
                    "modified": "付款应在验收合格后 30 日内支付。",
                    "paragraph_context": "补充约定：付款应在验收后支付。",
                }],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    inserted_text = "".join(document_xml.xpath(".//w:ins//w:t/text()", namespaces=W_NS))
    deleted_text = "".join(document_xml.xpath(".//w:del//w:delText/text()", namespaces=W_NS))
    assert inserted_text.count("付款应在验收合格后 30 日内支付。") == 1
    assert deleted_text.count("付款应在验收后支付。") == 1


def test_final_export_contains_no_revision_markup() -> None:
    modifications = [
        {
            "original": "合同份数：一式两份。",
            "modified": "合同份数：一式四份，甲乙双方各执两份。",
        }
    ]

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(modifications, ensure_ascii=False),
            "export_mode": "final",
        },
    )

    assert response.status_code == 200
    assert response.headers["content-disposition"] == 'attachment; filename="final_contract.docx"'
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    settings_xml = _read_docx_xml(response.content, "word/settings.xml")
    assert document_xml.find(".//w:ins", W_NS) is None
    assert document_xml.find(".//w:del", W_NS) is None
    assert settings_xml.find(".//w:trackRevisions", W_NS) is None
    assert "合同份数：一式四份，甲乙双方各执两份。" in "".join(
        document_xml.xpath(".//w:t/text()", namespaces=W_NS)
    )


def test_final_export_rejects_fuzzy_replacement_to_prevent_wrong_clause_edit() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [{"original": "合同份数一式两份", "modified": "合同份数：一式三份。"}],
                ensure_ascii=False,
            ),
            "export_mode": "final",
        },
    )

    assert response.status_code == 400
    assert "could not be located exactly" in response.json()["detail"]


def test_export_does_not_modify_same_text_in_header_or_footer() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_with_same_header_text(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [{"original": "合同份数：一式两份。", "modified": "合同份数：一式四份。"}],
                ensure_ascii=False,
            ),
            "export_mode": "final",
        },
    )

    assert response.status_code == 200
    header_xml = _read_docx_xml(response.content, "word/header1.xml")
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    assert "合同份数：一式两份。" in "".join(header_xml.xpath(".//w:t/text()", namespaces=W_NS))
    assert "合同份数：一式四份。" in "".join(document_xml.xpath(".//w:t/text()", namespaces=W_NS))


def test_final_export_rejects_missing_clause_without_exact_anchor() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [{"original": "【缺失该约定】", "modified": "新增通知条款。"}],
                ensure_ascii=False,
            ),
            "export_mode": "final",
        },
    )

    assert response.status_code == 400


def test_export_rejects_conflicting_replacements_for_one_original() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {"original": "合同份数：一式两份。", "modified": "合同份数：一式三份。"},
                    {"original": "合同份数：一式两份。", "modified": "合同份数：一式四份。"},
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 400
    assert "conflicting" in response.json()["detail"]


def test_export_inserts_missing_clause_after_anchor() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {
                        "original": "【缺失该约定】",
                        "modified": "新增税务条款：税费由乙方承担。",
                        "insert_after_text": "合同份数：一式两份。",
                    }
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    paragraphs = _paragraph_texts_from_xml(document_xml)

    assert paragraphs[0] == "合同份数：一式两份。"
    assert "新增税务条款：税费由乙方承担。" in paragraphs[1]
    assert document_xml.find(".//w:ins", W_NS) is not None


def test_export_rejects_fuzzy_insertion_anchor_to_preserve_precise_location() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {
                        "original": "【缺失该约定】",
                        "modified": "新增通知条款：双方应明确联系人与送达邮箱。",
                        "insert_after_text": "合同份数一式两份",
                    }
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 400
    assert "could not be located exactly" in response.json()["detail"]


def test_export_appends_missing_clause_modification() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [{"original": "【缺失该约定】", "modified": "新增条款：双方应明确通知联系人。"}],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    paragraphs_text = "\n".join(_paragraph_texts_from_xml(document_xml))

    assert "新增条款：双方应明确通知联系人。" in paragraphs_text


def test_export_allows_multiple_missing_clause_insertions() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {"original": "【缺失该约定】", "modified": "新增保密条款。", "insert_after_text": "合同份数：一式两份。"},
                    {"original": "【缺失该约定】", "modified": "新增数据安全条款。", "insert_after_text": "签订地点：未约定。"},
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200
    document_xml = _read_docx_xml(response.content, "word/document.xml")
    inserted_text = "".join(document_xml.xpath(".//w:ins//w:t/text()", namespaces=W_NS))
    assert "新增保密条款。" in inserted_text
    assert "新增数据安全条款。" in inserted_text


def test_export_matches_anchor_by_clause_heading() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_english_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {
                        "original": "【缺失该约定】",
                        "modified": "Counterparts. This Agreement may be executed in two counterparts.",
                        "insert_after_text": "Counterparts",
                    }
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    paragraphs = _paragraph_texts_from_xml(document_xml)

    assert paragraphs[1].startswith("(j) Counterparts.")
    assert paragraphs[2] == "Counterparts. This Agreement may be executed in two counterparts."


def test_export_replaces_text_by_clause_heading_similarity() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_english_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {
                        "original": "Counterparts. This Agreement may be executed in one or more counterparts",
                        "modified": "Counterparts. This Agreement may be executed in two counterparts.",
                    }
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    paragraphs_text = "\n".join(_paragraph_texts_from_xml(document_xml))

    assert "Counterparts. This Agreement may be executed in two counterparts." in paragraphs_text


def test_export_rejects_invalid_modifications_json() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"modifications": "{}"},
    )

    assert response.status_code == 400


def test_review_feedback_is_recorded(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("REVIEW_FEEDBACK_LOG", str(tmp_path / "feedback.jsonl"))
    response = client.post(
        "/api/review/feedback",
        json={
            "filename": "contract.docx",
            "risk_item": "付款与发票",
            "decision": "rejected",
            "note": "该条款已在附件中约定。",
        },
        headers={"X-Tenant-ID": "local"},
    )

    assert response.status_code == 201
    record = json.loads((tmp_path / "feedback.jsonl").read_text(encoding="utf-8"))
    assert record["decision"] == "rejected"
    assert record["tenant_id"] == "local"
    assert record["created_at"]
